import os
import json
import re
import phonenumbers
import subprocess
import time
from tkinter import Tk, Label, Entry, Button, Text, filedialog, messagebox, StringVar, OptionMenu
from validate_email import validate_email
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

# ------------------ Utility Functions ------------------

def validate_email_address(email):
    """Validate email using the validate_email library."""
    return validate_email(email)

def validate_phone_number(phone):
    """Validate phone number using the phonenumbers library."""
    try:
        phone_number = phonenumbers.parse(phone, "KE")
        return phonenumbers.is_valid_number(phone_number)
    except phonenumbers.NumberParseException:
        return False

def generate_unique_filename(base_name="cv"):
    """Generate a unique filename using a timestamp."""
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    return f"{base_name}_{timestamp}.docx"

def save_document_as_pdf(docx_filename):
    pdf_filename = docx_filename.replace(".docx", ".pdf")
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_filename], check=True)
    print(f"PDF saved as {pdf_filename}")
    return pdf_filename


def upload_to_google_drive(file_path):
    """Upload a file to Google Drive."""
    try:
        gauth = GoogleAuth()
        gauth.LocalWebserverAuth()  # Authenticate the user
        drive = GoogleDrive(gauth)

        file = drive.CreateFile({'title': os.path.basename(file_path)})
        file.SetContentFile(file_path)
        file.Upload()
        print(f"File uploaded successfully to Google Drive: {file['title']}")
        return True
    except Exception as e:
        print(f"Error uploading to Google Drive: {e}")
        return False

# ------------------ GUI Application ------------------
class CVGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("CV Generator")
        self.root.geometry("600x600")

        # Variables
        self.name_var = StringVar()
        self.phone_var = StringVar()
        self.email_var = StringVar()
        self.about_me_var = StringVar()
        self.template_var = StringVar(value="Template 1")

        # Widgets
        Label(root, text="CV Generator", font=("Arial", 20)).pack(pady=10)
        
        Label(root, text="Name:").pack()
        Entry(root, textvariable=self.name_var).pack()

        Label(root, text="Phone Number:").pack()
        Entry(root, textvariable=self.phone_var).pack()

        Label(root, text="Email:").pack()
        Entry(root, textvariable=self.email_var).pack()

        Label(root, text="About Me:").pack()
        self.about_me_text = Text(root, height=5, width=50)
        self.about_me_text.pack()

        Label(root, text="Choose Template:").pack()
        templates = ["Template 1", "Template 2", "Template 3"]
        OptionMenu(root, self.template_var, *templates).pack()

        Button(root, text="Generate CV", command=self.generate_cv).pack(pady=10)
        Button(root, text="Upload to Google Drive", command=self.upload_cv_to_drive).pack(pady=10)

        self.status_label = Label(root, text="", fg="green")
        self.status_label.pack(pady=10)

    def generate_cv(self):
        """Generate the CV based on user input and chosen template."""
        name = self.name_var.get().strip()
        phone = self.phone_var.get().strip()
        email = self.email_var.get().strip()
        about_me = self.about_me_text.get("1.0", "end").strip()
        template = self.template_var.get()

        if not name or not validate_phone_number(phone) or not validate_email_address(email):
            messagebox.showerror("Error", "Please enter valid Name, Phone, and Email.")
            return

        # Create document
        document = Document()
        document.add_heading(name, level=1)
        document.add_paragraph(f"Phone: {phone} | Email: {email}")
        document.add_heading("About Me", level=2)
        document.add_paragraph(about_me)

        # Add template-specific formatting
        if template == "Template 1":
            document.add_heading("Work Experience", level=2)
            document.add_paragraph("- Company ABC (2020-2023): Software Engineer")
        elif template == "Template 2":
            document.add_heading("Skills", level=2)
            document.add_paragraph("- Python, Java, Spring Boot")
        elif template == "Template 3":
            document.add_heading("Education", level=2)
            document.add_paragraph("- BSc Computer Science, XYZ University")

        # Save document
        filename = generate_unique_filename()
        document.save(filename)
        self.docx_filename = filename
        self.status_label.config(text=f"CV saved as {filename}")
        messagebox.showinfo("Success", f"CV generated successfully: {filename}")

        # Optionally convert to PDF
        save_as_pdf = messagebox.askyesno("Export", "Do you want to export the CV as a PDF?")
        if save_as_pdf:
            self.pdf_filename = save_document_as_pdf(filename)

    def upload_cv_to_drive(self):
        """Upload the generated CV to Google Drive."""
        if hasattr(self, 'pdf_filename') and os.path.exists(self.pdf_filename):
            success = upload_to_google_drive(self.pdf_filename)
            if success:
                messagebox.showinfo("Success", "CV uploaded to Google Drive successfully!")
            else:
                messagebox.showerror("Error", "Failed to upload CV to Google Drive.")
        else:
            messagebox.showerror("Error", "No PDF file found. Please generate a CV first.")

# ------------------ Main Function ------------------
if __name__ == "__main__":
    root = Tk()
    app = CVGeneratorApp(root)
    root.mainloop()
